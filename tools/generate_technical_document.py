from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "淘宝星河RPA数据采集与GMV报表系统技术开发文档.docx"


def set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_fill(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build() -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("淘宝星河 RPA 数据采集与 GMV 报表系统\n技术开发文档")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"版本：1.0    编制日期：{date.today().isoformat()}")

    document.add_heading("1. 项目概述", level=1)
    document.add_paragraph(
        "本项目通过 Python 与 Playwright 自动操作淘宝星河平台，完成项目数据采集、特殊报表下载、"
        "白名单数据匹配、GMV 汇总及 Excel 报表生成。除用户登录完成后的一次确认外，正式流程连续执行，"
        "无需额外人工确认。"
    )
    add_table(
        document,
        ["项目项", "说明"],
        [
            ["目标平台", "淘宝星河（adstar.alimama.com）"],
            ["主要技术", "Python 3.10、Playwright、pandas、openpyxl、PyYAML"],
            ["运行方式", "可视化 Chromium 浏览器 + 人工登录确认 + 自动业务流程"],
            ["采集周期", "2026-08-09 至本机运行当天；平台无当天数据时接受最近可用日期"],
            ["最终产物", "戈撒驰星河平台投流商家GMV报表（Excel）"],
        ],
    )

    document.add_heading("2. 系统架构", level=1)
    add_table(
        document,
        ["模块", "文件", "职责"],
        [
            ["程序入口", "main.py", "读取参数与配置，执行采集并触发最终报表生成"],
            ["配置加载", "rpa_collector/config.py", "读取并校验 YAML 配置"],
            ["RPA 采集", "rpa_collector/collector.py", "登录、导航、筛选、提取、下载及运行状态记录"],
            ["数据处理", "rpa_collector/reporting.py", "ID 白名单匹配、GMV 汇总及 Excel 输出"],
            ["业务配置", "config.yaml", "项目列表、日期、定位参数、下载与报表规则"],
        ],
    )

    document.add_heading("3. 端到端业务流程", level=1)
    steps = [
        "启动 Chromium 并打开淘宝星河；用户完成登录后在终端按一次 Enter。",
        "点击“我的星河”，依次处理四个标准项目。",
        "标准项目进入详情并点击“查看全部数据”，设置日期后提取商家 GMV；页面显示横线时内部记录为缺失值。",
        "进入“戈撒驰7月”项目及 gys7月达人卡片，选择推广效果、内容维度、日期和 30 天归因。",
        "生成报表，在下载记录中轮询最新任务，待“下载”按钮出现后保存 CSV 到项目目录。",
        "自动返回项目列表，进入 Y26-Gotyasatch-种草项目的“种草视频-第一波”详情。",
        "在推广效果汇总区设置日期和 30 天归因，等待 GMV 从切换前数值刷新后再提取。",
        "读取临时 GMV、七月 CSV 与字典工作簿，生成最终 Excel；任务完成后自动关闭浏览器。",
    ]
    for index, step in enumerate(steps, start=1):
        document.add_paragraph(f"{index}. {step}")

    document.add_heading("4. 项目规则", level=1)
    document.add_heading("4.1 标准 GMV 项目", level=2)
    add_bullets(
        document,
        [
            "Y26-Gotyasatch-大字报Q3",
            "Y26-Gotyasatch-UGCQ3",
            "Y26戈撒驰-UGC-Q3",
            "Y26-Gotyasatch-种草-Q3",
        ],
    )
    document.add_heading("4.2 七月项目", level=2)
    document.add_paragraph(
        "七月项目不直接采用页面汇总 GMV。系统仅读取字典工作簿的“CPUV底表”，先筛选时间不早于"
        " 2026-08-09 且推广目标等于 CPUV 的记录，再以“笔记/素材ID”为白名单，与下载 CSV 的“内容ID”"
        "进行标准化字符串匹配。只保留白名单命中的行，再对这些行的“商家GMV”求和。"
        "未命中的内容代表非本项目投流，必须排除。"
    )
    document.add_heading("4.3 种草视频特殊项目", level=2)
    document.add_paragraph(
        "目标卡片为 Y26-Gotyasatch戈撒驰书包-种草视频-第一波。日期使用可见的 .next-range-picker，"
        "通过起始日期/结束日期 placeholder 精确定位；归因口径限定在同一汇总筛选行内。切换为 30 天后，"
        "程序等待商家 GMV 发生变化，防止读取异步刷新前的 15 天旧值。"
    )

    document.add_heading("5. 关键技术实现", level=1)
    add_table(
        document,
        ["场景", "实现策略"],
        [
            ["项目详情入口", "以项目标题锚定 mux-card，悬停后点击 data-spm-click 包含 projectList_detail 的图标"],
            ["嵌套视频卡片", "精确匹配卡片标题，在卡片偏下位置悬停，再点击卡片内“查看详情”"],
            ["日期输入", "移除 readonly，仅对目标输入框 fill 完整日期并分别按 Enter；完成后二次读取校验"],
            ["平台日期滞后", "若结束日期被平台回退且不晚于请求日，则记录平台接受的最近可用日期"],
            ["金额解析", "识别人民币符号、千位分隔符与小数；横线转换为内部缺失标记"],
            ["下载任务", "定位最新操作时间记录并轮询，使用 Playwright 下载事件保存真实文件"],
            ["数据去重", "临时 GMV 按项目名更新，重跑不会为同一项目追加重复记录"],
            ["Excel 缺失值", "最终报表将 na、null 和空字符串统一输出为空白单元格"],
        ],
    )

    document.add_heading("6. 数据输入与输出", level=1)
    add_table(
        document,
        ["类型", "默认路径", "说明"],
        [
            ["临时 GMV", "runtime/extracted_data.json", "RPA 提取结果；不提交版本库"],
            ["七月原始报表", "output/july_reports/*.csv", "平台下载数据；不提交版本库"],
            ["CPUV 字典", "Gotyasatch戈撒驰书包日报0902.xlsx", "业务数据；不提交版本库"],
            ["最终报表", "output/final_reports/*.xlsx", "两张工作表：GMV报表、7月匹配明细"],
            ["运行日志", "collector.log、logs/", "调试与失败诊断；不提交版本库"],
        ],
    )
    document.add_paragraph(
        "最终主表第一行为“戈撒驰星河平台投流商家GMV报表”，第二行为“时间维度：2026-08-09 to current_day”，"
        "第三行为项目名和商家GMV表头。"
    )

    document.add_heading("7. 安装与运行", level=1)
    document.add_paragraph("在项目根目录执行：")
    for command in (
        "python -m venv .venv",
        ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt",
        ".\\.venv\\Scripts\\python.exe -m playwright install chromium",
        ".\\.venv\\Scripts\\python.exe main.py",
    ):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    document.add_paragraph(
        "浏览器打开后完成登录，在终端按 Enter。之后系统自动执行到结束。登录状态保存在 runtime/auth_state.json，"
        "过期时重新登录即可。"
    )

    document.add_heading("8. 配置说明", level=1)
    add_bullets(
        document,
        [
            "site.post_login_actions：控制页面动作与三类业务流程的执行顺序。",
            "browser.headless：是否隐藏浏览器；当前为 false，便于登录与排障。",
            "browser.pause_after_run：正式流程为 false，任务结束自动关闭。",
            "reporting.dictionary_sheets：固定为 CPUV底表。",
            "reporting.project_order：控制最终 Excel 项目顺序。",
            "reporting.start_date：控制报表起始日期，当前为 2026-08-09。",
        ],
    )

    document.add_heading("9. 异常处理与安全", level=1)
    add_bullets(
        document,
        [
            "选择器、日期确认、下载超时、金额解析或字典字段异常时立即抛出明确错误并记录日志。",
            "下载文件保存后检查文件存在且大小非零。",
            "最终 Excel 被打开锁定时，自动生成带时间后缀的新文件，避免覆盖失败。",
            "runtime、output、logs、登录状态、CSV、业务字典等数据文件均不提交 Git。",
            "不得绕过验证码或访问控制；采集范围应符合账号权限及平台规则。",
        ],
    )

    document.add_heading("10. 验收结果", level=1)
    add_bullets(
        document,
        [
            "六个目标项目均已完成实际页面导航验证。",
            "四个标准项目 GMV 提取通过，缺失值识别通过。",
            "七月内容维度报表生成、下载、CPUV底表匹配及 GMV 汇总通过。",
            "特殊视频项目顶部日期、30天归因与异步 GMV 刷新等待通过。",
            "完整端到端流程退出码为 0，登录确认后无额外人工暂停。",
            "最终 Excel 标题、时间维度、项目顺序、金额格式和空白缺失值均已验证。",
        ],
    )

    document.add_heading("11. 维护建议", level=1)
    add_bullets(
        document,
        [
            "平台页面升级后优先维护业务语义属性、文本锚点和模块作用域，避免依赖动态样式类。",
            "更换日报文件时同步更新 config.yaml 中 dictionary_file，并确认 CPUV底表及笔记/素材ID列存在。",
            "定期清理 output、logs 与 runtime 中的历史文件，但保留必要审计备份。",
            "正式交付前可使用 PyInstaller --onedir 打包，并一并配置 Playwright Chromium。",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
